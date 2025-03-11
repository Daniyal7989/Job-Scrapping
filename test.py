# import re
# import pdfplumber
# import streamlit as st
# from docx import Document
# import time
# from bs4 import BeautifulSoup
# from selenium import webdriver
# from selenium.webdriver.chrome.options import Options as ChromeOptions
# from selenium.webdriver.firefox.options import Options as FirefoxOptions
# from selenium.webdriver.chrome.service import Service as ChromeService
# from selenium.webdriver.firefox.service import Service as FirefoxService
# from webdriver_manager.chrome import ChromeDriverManager
# from webdriver_manager.firefox import GeckoDriverManager
# import pandas as pd
# import os

# # ---------------- Resume Parsing Functions ----------------
# def parse_resume(file):
#     file_extension = file.name.split('.')[-1].lower()
#     if file_extension == 'pdf':
#         text = extract_text_from_pdf(file)
#     elif file_extension == 'docx':
#         text = extract_text_from_docx(file)
#     else:
#         raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
    
#     parsed_data = {
#         "name": extract_name(text),
#         "email": extract_email(text),
#         "phone": extract_phone(text),
#         "location": extract_location(text),
#         "skills": extract_skills(text),
#         "languages": extract_languages(text),
#         "experience": extract_experience(text),
#         "qualifications": extract_qualifications(text),
#         "employment_history": extract_employment_history(text),
#         "profile_summary": extract_profile_summary(text)
#     }
#     return parsed_data

# def extract_text_from_pdf(file):
#     try:
#         with pdfplumber.open(file) as pdf:
#             text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#         return text.strip()
#     except Exception as e:
#         raise ValueError(f"Error extracting text from PDF: {e}")

# def extract_text_from_docx(file):
#     try:
#         doc = Document(file)
#         text = "\n".join([para.text for para in doc.paragraphs])
#         return text.strip()
#     except Exception as e:
#         raise ValueError(f"Error extracting text from DOCX: {e}")

# def extract_name(text):
#     lines = text.split("\n")
#     possible_names = [line for line in lines[:5] if len(line.split()) in [2, 3]]
#     return possible_names[0] if possible_names else "Unknown"

# def extract_email(text):
#     match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
#     return match.group(0) if match else "Not Found"

# def extract_phone(text):
#     match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
#     return match.group(0) if match else "Not Found"

# def extract_location(text):
#     lines = text.split("\n")
#     location_section_start = 0
#     for i, line in enumerate(lines):
#         if "Details" in line:
#             location_section_start = i
#             break
#     for i in range(location_section_start, min(location_section_start + 10, len(lines))):
#         if lines[i] in ["United States", "USA", "U.S.", "U.S.A."]:
#             city_line = i - 1 if i > 0 else 0
#             return f"{lines[city_line]}, {lines[i]}"
#     city_state_pattern = re.search(r"([A-Za-z\s]+),?\s+([A-Za-z\s]+)", text[:500])
#     if city_state_pattern:
#         return city_state_pattern.group(0)
#     return "Location Not Found"

# def extract_skills(text):
#     skill_keywords = [
#         "Python", "Flask", "Django", "Machine Learning", "SQL", "Java", "React", "AWS",
#         "JavaScript", "HTML", "CSS", "TensorFlow", "Pandas", "NumPy", "Docker", "Kubernetes",
#         "Git", "Azure", "Linux", "Node.js", "C#", "C++", "Go", "PHP", "TypeScript", "Tableau", 
#         "Power BI", "Jupyter", "Spark", "Hadoop", "Scala", "Cloud", "Vagrant", "LLMs", "GPT",
#         "Re-enforcement", "Site Reliability", "DevOps", "Microservices", "NOSQL", 
#         "Apache Kafka", "Apache Webserver", "Blockchain", "Performance engineering",
#         "AI model Training", "Data Science", "Feature Engineering", "AI", "Shell Script", 
#         "Intrusion Detection", "Matlab", "R", "Agile", "SDLC"
#     ]
#     skills_section_match = re.search(r"Skills\n(.*?)(?:\n\n|\nProfile|\nEmployment)", text, re.DOTALL)
#     skills_text = skills_section_match.group(1) if skills_section_match else text
#     found_skills = [skill for skill in skill_keywords if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
#     additional_skills = [line.strip() for line in skills_text.split("\n") if line.strip() and len(line.strip()) > 2]
#     all_skills = list(set(found_skills + additional_skills))
#     return all_skills

# def extract_languages(text):
#     languages_section = re.search(r"Languages\n(.*?)(?:\n\n|\nEducation|\nEmployment)", text, re.DOTALL)
#     if languages_section:
#         languages_text = languages_section.group(1)
#         languages = [lang.strip() for lang in languages_text.split("\n") if lang.strip()]
#         return languages
#     common_languages = ["English", "Spanish", "French", "German", "Chinese", "Japanese", 
#                         "Arabic", "Hindi", "Bengali", "Russian", "Portuguese"]
#     found_languages = [lang for lang in common_languages if re.search(rf"\b{lang}\b", text)]
#     return found_languages if found_languages else ["Not Found"]

# def extract_experience(text):
#     years_of_experience = re.findall(r"(\d{1,2})\+?\s*years?(?:\s+of\s+experience)?", text)
#     if years_of_experience:
#         max_experience = max(map(int, years_of_experience))
#         return f"{max_experience} years of experience"
#     employment_pattern = re.findall(r"(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", text, re.IGNORECASE)
#     if employment_pattern:
#         return f"{len(employment_pattern)} employment periods found"
#     return "Experience not found"

# def extract_qualifications(text):
#     qualification_keywords = [
#         "Bachelor's", "Master's", "PhD", "Degree", "Certification", "Diploma", "BSc", "MSc", 
#         "B.A.", "M.A.", "BTech", "MTech", "MBA", "Engineering", "Architecture", "Computer Science", 
#         "Information Technology", "Data Science", "Machine Learning", "AI", "Software Engineering",
#         "CSSA", "Certified", "Business", "MIS", "University"
#     ]
#     education_section = re.search(r"Education\n(.*?)(?:\n\n|$)", text, re.DOTALL)
#     if education_section:
#         edu_text = education_section.group(1)
#         education_entries = [line.strip() for line in edu_text.split("\n") if line.strip() 
#                             and any(keyword in line for keyword in ["Bachelor", "Master", "PhD", "Degree"])]
#         if education_entries:
#             return education_entries
#     found_qualifications = [qual for qual in qualification_keywords if re.search(rf"\b{qual}\b", text, re.IGNORECASE)]
#     return list(set(found_qualifications)) if found_qualifications else ["Not Found"]

# def extract_employment_history(text):
#     employment_section = re.search(r"Employment History\n(.*?)(?:\nEducation|\n\nEducation|$)", text, re.DOTALL | re.IGNORECASE)
#     if not employment_section:
#         return ["Employment history not found"]
#     employment_text = employment_section.group(1)
#     job_pattern = re.findall(r"(.*?),\s+(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
#                              employment_text, re.IGNORECASE | re.MULTILINE)
#     if not job_pattern:
#         job_pattern = re.findall(r"(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
#                                 employment_text, re.IGNORECASE | re.MULTILINE)
#     if not job_pattern:
#         lines = employment_text.split("\n")
#         jobs = []
#         for i, line in enumerate(lines):
#             if re.search(r"\b(architect|analyst|lead|principal|manager|director|engineer)\b", line, re.IGNORECASE):
#                 if i < len(lines) - 1 and re.search(r"\d{4}", lines[i+1]):
#                     jobs.append(line)
#         return jobs if jobs else ["Could not parse employment details"]
    
#     formatted_jobs = []
#     for job in job_pattern:
#         if len(job) == 5:
#             formatted_jobs.append(f"{job[0]} at {job[1]}, {job[2]} ({job[3]} - {job[4]})")
#         elif len(job) == 4:
#             formatted_jobs.append(f"{job[0]} at {job[1]} ({job[2]} - {job[3]})")
#     return formatted_jobs

# def extract_profile_summary(text):
#     profile_section = re.search(r"Profile\n(.*?)(?:\nEmployment|\n\nEmployment)", text, re.DOTALL)
#     if profile_section:
#         profile_text = profile_section.group(1).strip()
#         if len(profile_text) > 500:
#             return profile_text[:497] + "..."
#         return profile_text
#     alt_section = re.search(r"(Summary|About)\n(.*?)(?:\n\n|\nSkills|\nExperience)", text, re.DOTALL)
#     if alt_section:
#         return alt_section.group(2).strip()
#     return "Profile summary not found"



# # ---------------- Job Scraping and Matching Functions ----------------

# def get_chrome_binary():
#     possible_paths = [
#         "/opt/google/chrome/chrome",
#         "/usr/bin/chromium",
#         "/usr/bin/chromium-browser"
#     ]
#     for path in possible_paths:
#         if os.path.exists(path):
#             return path
#     return None

# def init_chrome_driver():
#     options = ChromeOptions()
#     options.add_argument("--headless")
#     options.add_argument("--disable-gpu")
#     options.add_argument("--no-sandbox")
#     options.add_argument("--disable-dev-shm-usage")
#     options.add_argument("--remote-debugging-port=9222")
#     options.add_argument("--start-maximized")
#     options.add_argument("--disable-blink-features=AutomationControlled")
#     options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64)")
#     chrome_binary = get_chrome_binary()
#     if not chrome_binary:
#         raise Exception("No Chrome or Chromium binary found!")
#     options.binary_location = chrome_binary
#     if "google/chrome" in chrome_binary:
#         driver_version = "133.0.6943.126"
#     else:
#         driver_version = "120.0.6099.224"
#     driver_path = ChromeDriverManager(driver_version=driver_version).install()
#     os.chmod(driver_path, 0o755)
#     service = ChromeService(driver_path)
#     driver = webdriver.Chrome(service=service, options=options)
#     return driver

# def init_firefox_driver():
#     options = FirefoxOptions()
#     options.add_argument("--headless")
#     driver_path = GeckoDriverManager().install()
#     service = FirefoxService(driver_path)
#     driver = webdriver.Firefox(service=service, options=options)
#     return driver

# def init_driver():
#     try:
#         driver = init_chrome_driver()
#         print("Initialized Chrome/Chromium driver.")
#         return driver
#     except Exception as e:
#         print("Chrome driver initialization failed, trying Firefox. Error:", e)
#         driver = init_firefox_driver()
#         print("Initialized Firefox driver.")
#         return driver

# def scrape_dice(job_role, driver, preferred_location="", preferred_salary=""):
#     print(f"🔍 Scraping Dice for: {job_role}")
#     formatted_role = job_role.replace(" ", "%20")
#     # Use location and salary from user preferences if provided
#     loc_param = preferred_location if preferred_location and preferred_location != "Select Location" else ""
#     salary_param = preferred_salary if preferred_salary and preferred_salary != "Select Salary Range" else ""
#     url = f"https://www.dice.com/jobs?q={formatted_role}&location={loc_param}&salary={salary_param}&countryCode=US&page=1&pageSize=20&language=en"
#     driver.get(url)
#     time.sleep(5)
#     soup = BeautifulSoup(driver.page_source, "html.parser")
#     job_elements = soup.find_all("div", class_="card")
#     jobs = []
#     for job in job_elements:
#         title_element = job.find("h5")
#         link_element = job.find("a", href=True)
#         location_element = job.find("span", attrs={"data-cy": "search-result-location"})
#         company_name = job.find("a", attrs={"aria-label": lambda x: x and "company page for" in x})
#         job_type_element = job.find("span", attrs={"data-cy": "search-result-employment-type"})
#         posted_time_element = job.find("span", attrs={"data-cy": "card-posted-date"})
#         updated_time_element = job.find("span", attrs={"data-cy": "card-modified-date"})
#         description_element = job.find("div", attrs={"data-cy": "card-summary"})
#         if title_element and link_element:
#             jobs.append({
#                 "Job Role": job_role,
#                 "Title": title_element.text.strip(),
#                 "Apply Link": link_element["href"],
#                 "Location": location_element.text.strip() if location_element else "N/A",
#                 "Job Type": job_type_element.text.strip() if job_type_element else "N/A",
#                 "Posted Time": posted_time_element.text.strip() if posted_time_element else "N/A",
#                 "Updated Time": updated_time_element.text.strip() if updated_time_element else "N/A",
#                 "Description": description_element.text.strip() if description_element else "N/A",
#                 "Company Name": company_name.text.strip() if company_name else "N/A"
#             })
#     return jobs

# def extract_career_goal_keywords(career_goals):
#     if not career_goals:
#         return {}
#     goals_data = {
#         "job_titles": [],
#         "locations": [],
#         "work_preferences": [],
#         "industries": []
#     }
#     job_title_patterns = [
#         r"(?:seeking|looking for|interested in)(?:\s+a)?\s+(?:role|position|job|career)?\s+(?:as|in)?\s+(?:an?)?\s+((?:senior|junior|lead|principal|staff)?\s*[a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))",
#         r"(?:senior|junior|lead|principal|staff)?\s*([a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))\s+(?:role|position|job)",
#         r"(?:become|be)\s+(?:a|an)\s+((?:senior|junior|lead|principal|staff)?\s*[a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))"
#     ]
#     for pattern in job_title_patterns:
#         matches = re.findall(pattern, career_goals.lower())
#         if matches:
#             goals_data["job_titles"].extend([title.strip().title() for title in matches])
#     common_titles = [
#         "software engineer", "data scientist", "machine learning engineer", 
#         "devops engineer", "full stack developer", "frontend developer", 
#         "backend developer", "site reliability engineer", "cloud engineer",
#         "ai engineer", "web developer", "product manager", "project manager",
#         "ux designer", "systems engineer", "solutions architect",
#         "data engineer", "ml engineer", "research scientist"
#     ]
#     for title in common_titles:
#         if title in career_goals.lower():
#             goals_data["job_titles"].append(title.title())
#     location_patterns = [
#         r"(?:in|near|around|based in)\s+([a-z\s]+(?:area|city|region|county))",
#         r"(?:in|near|around|based in)\s+([a-z]+,\s*[a-z]{2})",
#         r"(?:in|near|around|based in)\s+([a-z\s]+)"
#     ]
#     for pattern in location_patterns:
#         matches = re.findall(pattern, career_goals.lower())
#         if matches:
#             goals_data["locations"].extend([loc.strip().title() for loc in matches])
#     common_cities = ["new york", "san francisco", "seattle", "boston", "austin", "chicago", 
#                      "los angeles", "denver", "washington dc", "atlanta", "dallas"]
#     for city in common_cities:
#         if city in career_goals.lower():
#             goals_data["locations"].append(city.title())
#     preferences = ["remote", "hybrid", "onsite", "in-office", "work from home", "flexible", "part-time", "full-time"]
#     for pref in preferences:
#         if pref in career_goals.lower():
#             goals_data["work_preferences"].append(pref.title())
#     industries = ["tech", "healthcare", "finance", "fintech", "edtech", "e-commerce", "retail", 
#                  "gaming", "entertainment", "education", "manufacturing", "automotive", 
#                  "energy", "sustainability", "non-profit", "government", "consulting"]
#     for industry in industries:
#         if industry in career_goals.lower():
#             goals_data["industries"].append(industry.title())
#     return goals_data

# def get_relevant_job_titles(parsed_data):
#     """Determine job titles based on career goals and resume details.
#        Priority is given if the career goal keywords appear in the resume skills."""
#     job_titles = []
#     career_goals = parsed_data.get('career_goals', "")
#     career_keywords = extract_career_goal_keywords(career_goals)
#     career_job_titles = career_keywords.get("job_titles", [])
#     # Check if any career title keyword appears in the parsed skills
#     valid_titles = []
#     for title in career_job_titles:
#         for skill in parsed_data.get('skills', []):
#             if title.lower() in skill.lower():
#                 valid_titles.append(title)
#                 break
#     if valid_titles:
#         job_titles.extend(valid_titles)
#     else:
#         # If not, use career goals job titles (first priority) even if not in skills
#         if career_job_titles:
#             job_titles.extend(career_job_titles)
#     # Fall back to employment history if needed
#     if len(job_titles) < 3:
#         for job in parsed_data.get('employment_history', []):
#             if "at" in job:
#                 title = job.split("at")[0].strip().rstrip(",")
#                 job_titles.append(title)
#         key_skills = []
#         priority_skills = ["AI", "Machine Learning", "Python", "Java", "Cloud", "DevOps", "Architect", 
#                            "Data Science", "Full Stack", "Backend", "Frontend", "Software Engineer"]
#         for skill in priority_skills:
#             if any(skill.lower() in s.lower() for s in parsed_data.get('skills', [])):
#                 key_skills.append(skill)
#         for skill in key_skills[:3]:
#             job_titles.append(f"{skill} Engineer")
#             job_titles.append(f"{skill} Developer")
#     unique_titles = list(set(job_titles))
#     return unique_titles[:5]


# def calculate_job_match_score(job, parsed_data, motivation_factors):
#     """Calculate a job match score based on resume data and motivation factors."""
#     total_score = 0
#     max_possible_score = sum([factor['weight'] for factor in motivation_factors])
#     factor_scores = {}
    
#     for factor in motivation_factors:
#         factor_name = factor['name']
#         factor_weight = factor['weight']
#         score = 0
        
#         if factor_name == "Salary":
#             salary_keywords = ["competitive salary", "excellent compensation", "high pay", "top pay"]
#             if any(keyword in job['Description'].lower() for keyword in salary_keywords):
#                 score = factor_weight
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Location":
#             user_location = parsed_data.get('preferred_location', parsed_data.get('location', '')).lower()
#             if user_location and user_location != "location not found":
#                 if user_location in job['Location'].lower():
#                     score = factor_weight
#                 elif "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
#                     score = factor_weight * 0.8
#                 else:
#                     score = 0
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Company Culture":
#             culture_keywords = ["work-life balance", "inclusive", "diversity", "collaborative", "flexible", "innovative", "team-oriented", "growth opportunities"]
#             matches = sum(1 for keyword in culture_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 3, 1.0)
#         elif factor_name == "Work-Life Balance":
#             balance_keywords = ["flexible hours", "work-life balance", "remote work", "flexible schedule", "family friendly", "paid time off", "pto", "vacation"]
#             matches = sum(1 for keyword in balance_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Growth Opportunities":
#             growth_keywords = ["career advancement", "professional development", "promotion", "training", "learning", "mentorship", "growth", "career path"]
#             matches = sum(1 for keyword in growth_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Job Stability":
#             stability_keywords = ["established company", "long-term", "stable", "security", "permanent position", "industry leader"]
#             matches = sum(1 for keyword in stability_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Remote Work":
#             if "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
#                 score = factor_weight
#             elif "hybrid" in job['Job Type'].lower() or "hybrid" in job['Description'].lower():
#                 score = factor_weight * 0.7
#             else:
#                 score = 0
#         elif factor_name == "Company Size":
#             score = factor_weight * 0.5
#         elif factor_name == "Industry Match":
#             if 'career_goals' in parsed_data and parsed_data['career_goals']:
#                 career_goal_keywords = extract_career_goal_keywords(parsed_data['career_goals'])
#                 industries = career_goal_keywords.get("industries", [])
#                 if industries:
#                     if any(industry.lower() in job['Description'].lower() for industry in industries):
#                         score = factor_weight
#                     else:
#                         score = 0
#                 else:
#                     score = factor_weight * 0.5
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Skill Match":
#             if parsed_data.get('skills'):
#                 matches = sum(1 for skill in parsed_data['skills'] if skill.lower() in job['Description'].lower())
#                 score = factor_weight * min(matches / max(5, len(parsed_data['skills']) * 0.3), 1.0)
#             else:
#                 score = factor_weight * 0.5
#         else:
#             score = factor_weight * 0.5
        
#         factor_scores[factor_name] = score
#         total_score += score

#     percentage_score = (total_score / max_possible_score) * 100 if max_possible_score > 0 else 50

#     # Map the percentage score into three categories: High, Medium, Low
#     if percentage_score >= 80:
#         match_category = "High"
#     elif percentage_score >= 50:
#         match_category = "Medium"
#     else:
#         match_category = "Low"

#     return {
#         "total_score": total_score,
#         "percentage_score": percentage_score,
#         "match_category": match_category,
#         "factor_scores": factor_scores
#     }


# # ---------------- Main Streamlit Application ----------------

# def calculate_job_match_score(job, parsed_data, motivation_factors):
#     """Calculate a job match score based on resume data and motivation factors."""
#     total_score = 0
#     max_possible_score = sum([factor['weight'] for factor in motivation_factors])
#     factor_scores = {}
    
#     for factor in motivation_factors:
#         factor_name = factor['name']
#         factor_weight = factor['weight']
#         score = 0
        
#         if factor_name == "Salary":
#             salary_keywords = ["competitive salary", "excellent compensation", "high pay", "top pay"]
#             if any(keyword in job['Description'].lower() for keyword in salary_keywords):
#                 score = factor_weight
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Location":
#             user_location = parsed_data.get('preferred_location', parsed_data.get('location', '')).lower()
#             if user_location and user_location != "location not found":
#                 if user_location in job['Location'].lower():
#                     score = factor_weight
#                 elif "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
#                     score = factor_weight * 0.8
#                 else:
#                     score = 0
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Company Culture":
#             culture_keywords = ["work-life balance", "inclusive", "diversity", "collaborative", "flexible", "innovative", "team-oriented", "growth opportunities"]
#             matches = sum(1 for keyword in culture_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 3, 1.0)
#         elif factor_name == "Work-Life Balance":
#             balance_keywords = ["flexible hours", "work-life balance", "remote work", "flexible schedule", "family friendly", "paid time off", "pto", "vacation"]
#             matches = sum(1 for keyword in balance_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Growth Opportunities":
#             growth_keywords = ["career advancement", "professional development", "promotion", "training", "learning", "mentorship", "growth", "career path"]
#             matches = sum(1 for keyword in growth_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Job Stability":
#             stability_keywords = ["established company", "long-term", "stable", "security", "permanent position", "industry leader"]
#             matches = sum(1 for keyword in stability_keywords if keyword in job['Description'].lower())
#             score = factor_weight * min(matches / 2, 1.0)
#         elif factor_name == "Remote Work":
#             if "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
#                 score = factor_weight
#             elif "hybrid" in job['Job Type'].lower() or "hybrid" in job['Description'].lower():
#                 score = factor_weight * 0.7
#             else:
#                 score = 0
#         elif factor_name == "Company Size":
#             score = factor_weight * 0.5
#         elif factor_name == "Industry Match":
#             if 'career_goals' in parsed_data and parsed_data['career_goals']:
#                 career_goal_keywords = extract_career_goal_keywords(parsed_data['career_goals'])
#                 industries = career_goal_keywords.get("industries", [])
#                 if industries:
#                     if any(industry.lower() in job['Description'].lower() for industry in industries):
#                         score = factor_weight
#                     else:
#                         score = 0
#                 else:
#                     score = factor_weight * 0.5
#             else:
#                 score = factor_weight * 0.5
#         elif factor_name == "Skill Match":
#             if parsed_data.get('skills'):
#                 matches = sum(1 for skill in parsed_data['skills'] if skill.lower() in job['Description'].lower())
#                 score = factor_weight * min(matches / max(5, len(parsed_data['skills']) * 0.3), 1.0)
#             else:
#                 score = factor_weight * 0.5
#         else:
#             score = factor_weight * 0.5
        
#         factor_scores[factor_name] = score
#         total_score += score

#     percentage_score = (total_score / max_possible_score) * 100 if max_possible_score > 0 else 50

#     # Map the percentage score into three categories: High, Medium, Low
#     if percentage_score >= 80:
#         match_category = "High"
#     elif percentage_score >= 50:
#         match_category = "Medium"
#     else:
#         match_category = "Low"

#     return {
#         "total_score": total_score,
#         "percentage_score": percentage_score,
#         "match_category": match_category,
#         "factor_scores": factor_scores
#     }

# # ---------------- Initialize Session State Variables (only once) ----------------
# if 'current_step' not in st.session_state:
#     st.session_state.current_step = 1
# if 'parsed_data' not in st.session_state:
#     st.session_state.parsed_data = None
# if 'career_goals' not in st.session_state:
#     st.session_state.career_goals = ""
# if 'preferred_location' not in st.session_state:
#     st.session_state.preferred_location = "Select Location"
# if 'preferred_salary' not in st.session_state:
#     st.session_state.preferred_salary = "Select Salary Range"
# if 'motivation_factors' not in st.session_state:
#     st.session_state.motivation_factors = [
#         {"name": "Salary", "weight": 5},
#         {"name": "Location", "weight": 4},
#         {"name": "Company Culture", "weight": 3},
#         {"name": "Work-Life Balance", "weight": 3},
#         {"name": "Growth Opportunities", "weight": 4},
#         {"name": "Skill Match", "weight": 5}
#     ]
# if 'job_results' not in st.session_state:
#     st.session_state.job_results = None

# # ---------------- Main Function ----------------
# def main():
#     st.title("Enhanced Resume Parser with Job Matching")
#     st.write("Follow the steps below to upload your resume and find matching job opportunities.")

#     # ---------- STEP 1: Upload and Parse Resume ----------
#     if st.session_state.current_step == 1:
#         st.header("Step 1: Upload Your Resume")
#         if st.session_state.parsed_data is None:
#             uploaded_file = st.file_uploader("Choose a PDF or DOCX resume", type=["pdf", "docx"])
#             if uploaded_file is None:
#                 st.error("Please upload your CV to continue.")
#             else:
#                 if st.button("Parse Resume", key="parse_resume_btn"):
#                     with st.spinner("Parsing Resume..."):
#                         try:
#                             parsed_data = parse_resume(uploaded_file)
#                             st.session_state.parsed_data = parsed_data
#                             st.success("Resume parsed successfully!")
#                         except Exception as e:
#                             st.error(f"Error parsing resume: {e}")
#         else:
#             st.subheader("Parsed Resume Details")
#             display_parsed_resume(st.session_state.parsed_data)
#         # Navigation Buttons for Step 1
#         col1, col2 = st.columns(2)
#         with col1:
#             st.write("")  # Placeholder for alignment
#         with col2:
#             if st.button("Next", key="next_after_parse"):
#                 st.session_state.current_step = 2

#     # ---------- STEP 2: Career Goals and Preferences ----------
#     elif st.session_state.current_step == 2:
#         st.header("Step 2: Define Your Career Purpose and Preferences")
#         career_goals = st.text_area(
#             "Career Goals and Preferences",
#             value=st.session_state.career_goals,
#             height=200,
#             placeholder="Example: I am seeking a Senior Machine Learning Engineer role in the healthcare industry..."
#         )
#         location = st.selectbox(
#             "Select Preferred Job Location", 
#             ["Select Location", "New York", "San Francisco", "Los Angeles", "Boston", "Chicago", "Remote"],
#             index=0
#         )
#         salary_range = st.selectbox(
#             "Select Desired Salary Range", 
#             ["Select Salary Range", "Below $50k", "$50k-$70k", "$70k-$90k", "$90k-$110k", "Above $110k"],
#             index=0
#         )
#         # Navigation Buttons for Step 2
#         col1, col2, col3 = st.columns(3)
#         with col1:
#             if st.button("← Back", key="back_to_step1"):
#                 st.session_state.current_step = 1
#         with col2:
#             st.write("")  # Spacer
#         with col3:
#             if st.button("Next", key="next_step2"):
#                 if not career_goals.strip():
#                     st.error("Career goals and preferences cannot be empty!")
#                 elif location == "Select Location":
#                     st.error("Please select a job location.")
#                 elif salary_range == "Select Salary Range":
#                     st.error("Please select a desired salary range.")
#                 else:
#                     st.session_state.career_goals = career_goals
#                     st.session_state.preferred_location = location
#                     st.session_state.preferred_salary = salary_range
#                     if st.session_state.parsed_data:
#                         st.session_state.parsed_data['career_goals'] = career_goals
#                         st.session_state.parsed_data['preferred_location'] = location
#                         st.session_state.parsed_data['preferred_salary'] = salary_range
#                     st.success("Career goals and preferences saved!")
#                     if st.button("Next", key="next_after_step2"):
#                         st.session_state.current_step = 3

#     # ---------- STEP 3: Motivation Matrix ----------
#     elif st.session_state.current_step == 3:
#         st.header("Step 3: Customize Your Motivation Matrix")
#         updated_factors = []
#         for factor in st.session_state.motivation_factors:
#             weight = st.slider(
#                 f"{factor['name']}",
#                 min_value=1,
#                 max_value=10,
#                 value=factor['weight'],
#                 help=f"Set importance of {factor['name']}",
#                 key=f"slider_{factor['name']}"
#             )
#             updated_factors.append({"name": factor['name'], "weight": weight})
#         st.session_state.motivation_factors = updated_factors
#         # Navigation Buttons for Step 3
#         col1, col2 = st.columns(2)
#         with col1:
#             if st.button("← Back", key="back_to_step2"):
#                 st.session_state.current_step = 2
#         with col2:
#             if st.button("Next", key="next_step3"):
#                 st.success("Motivation matrix saved!")
#                 st.session_state.current_step = 4

#     # ---------- STEP 4: Job Matching and Filtering ----------
#     elif st.session_state.current_step == 4:
#         st.header("Step 4: Job Matching Results")
#         if st.session_state.parsed_data is None:
#             st.warning("Please upload your resume first.")
#             st.session_state.current_step = 1
#         else:
#             col1, col2 = st.columns(2)
#             with col1:
#                 if st.button("← Back", key="back_to_step3"):
#                     st.session_state.current_step = 3
#             with col2:
#                 if st.button("Find Matching Jobs", key="find_jobs_btn"):
#                     with st.spinner("Searching for matching jobs..."):
#                         try:
#                             job_titles = get_relevant_job_titles(st.session_state.parsed_data)
#                             st.write("Searching for the following job roles:")
#                             for title in job_titles:
#                                 st.write(f"- {title}")
#                             driver = init_driver()
#                             all_jobs = []
#                             for job_title in job_titles:
#                                 jobs = scrape_dice(job_title, driver, st.session_state.preferred_location, st.session_state.preferred_salary)
#                                 all_jobs.extend(jobs)
#                             driver.quit()
#                             scored_jobs = []
#                             for job in all_jobs:
#                                 match_score = calculate_job_match_score(
#                                     job,
#                                     st.session_state.parsed_data,
#                                     st.session_state.motivation_factors
#                                 )
#                                 job_with_score = {**job, "match_score": match_score}
#                                 scored_jobs.append(job_with_score)
#                             sorted_jobs = sorted(scored_jobs, key=lambda x: x["match_score"]["percentage_score"], reverse=True)
#                             st.session_state.job_results = sorted_jobs
#                             st.success(f"Found {len(sorted_jobs)} matching jobs!")
#                         except Exception as e:
#                             st.error(f"Error finding matching jobs: {e}")
#             if st.session_state.job_results is not None:
#                 if st.button("Next", key="next_step4"):
#                     st.session_state.current_step = 5

#     # ---------- STEP 5: Detailed Summary Report with CSV Download ----------
#     elif st.session_state.current_step == 5:
#         st.header("Step 5: Detailed Summary Report")
#         if st.session_state.job_results is None:
#             st.warning("No job results available. Please run a job search first.")
#             st.session_state.current_step = 4
#         else:
#             # Group jobs by match category
#             report = {"High": [], "Medium": [], "Low": []}
#             for job in st.session_state.job_results:
#                 cat = job["match_score"]["match_category"]
#                 report.setdefault(cat, []).append(job)
            
#             st.markdown("### Summary Counts")
#             st.write(f"**High:** {len(report.get('High', []))}  |  **Medium:** {len(report.get('Medium', []))}  |  **Low:** {len(report.get('Low', []))}")
#             st.markdown("---")
#             # Prepare a DataFrame for CSV download (all job results)
#             jobs_df = pd.DataFrame([
#                 {
#                     "Title": job["Title"],
#                     "Company": job["Company Name"],
#                     "Location": job["Location"],
#                     "Job Type": job["Job Type"],
#                     "Match Score": f"{job['match_score']['percentage_score']:.1f}%",
#                     "Match Category": job["match_score"]["match_category"],
#                     "Posted": job["Posted Time"],
#                     "Apply Link": job["Apply Link"]
#                 }
#                 for job in st.session_state.job_results
#             ])
#             st.download_button(
#                 label="Download Job Listings as CSV",
#                 data=jobs_df.to_csv(index=False).encode('utf-8'),
#                 file_name="job_matches.csv",
#                 mime="text/csv"
#             )
#             # Display job details grouped by category
#             for cat in ["High", "Medium", "Low"]:
#                 jobs_in_cat = report.get(cat, [])
#                 if jobs_in_cat:
#                     with st.expander(f"{cat} Match Jobs ({len(jobs_in_cat)})"):
#                         for job in jobs_in_cat:
#                             st.markdown(f"**{job['Title']} at {job['Company Name']}** - {job['match_score']['percentage_score']:.1f}% Match")
#                             job_col1, job_col2 = st.columns([2, 1])
#                             with job_col1:
#                                 st.write(f"**Company:** {job['Company Name']}")
#                                 st.write(f"**Location:** {job['Location']}")
#                                 st.write(f"**Job Type:** {job['Job Type']}")
#                                 st.write(f"**Posted:** {job['Posted Time']}")
#                                 st.write("**Description:**")
#                                 st.write(job['Description'])
#                                 st.markdown(f"[Apply Now]({job['Apply Link']})")
#                             with job_col2:
#                                 st.write("**Match Breakdown:**")
#                                 for factor, score in job["match_score"]["factor_scores"].items():
#                                     max_score = next((f["weight"] for f in st.session_state.motivation_factors if f["name"] == factor), 5)
#                                     st.write(f"- {factor}: {score}/{max_score}")
#                                 match_color = {"High": "green", "Medium": "orange", "Low": "red"}.get(job["match_score"]["match_category"], "gray")
#                                 st.markdown(
#                                     f"<div style='background-color:{match_color}; height:20px; width:{job['match_score']['percentage_score']}%; border-radius:5px;'></div>",
#                                     unsafe_allow_html=True
#                                 )
#                                 st.write(f"{job['match_score']['percentage_score']:.1f}% Match")
#                             st.markdown("---")
#             if st.button("← Back", key="back_to_step4"):
#                 st.session_state.current_step = 4

# if __name__ == "__main__":
#     main()

import re
import pdfplumber
import streamlit as st
from docx import Document
import time
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.firefox import GeckoDriverManager
import pandas as pd
import os

# ---------------- Resume Parsing Functions ----------------
def parse_resume(file):
    file_extension = file.name.split('.')[-1].lower()
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
    
    parsed_data = {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "location": extract_location(text),
        "skills": extract_skills(text),
        "languages": extract_languages(text),
        "experience": extract_experience(text),
        "qualifications": extract_qualifications(text),
        "employment_history": extract_employment_history(text),
        "profile_summary": extract_profile_summary(text)
    }
    return parsed_data

def extract_text_from_pdf(file):
    try:
        with pdfplumber.open(file) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {e}")

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {e}")

def extract_name(text):
    lines = text.split("\n")
    possible_names = [line for line in lines[:5] if len(line.split()) in [2, 3]]
    return possible_names[0] if possible_names else "Unknown"

def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else "Not Found"

def extract_phone(text):
    match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    return match.group(0) if match else "Not Found"

def extract_location(text):
    lines = text.split("\n")
    location_section_start = 0
    for i, line in enumerate(lines):
        if "Details" in line:
            location_section_start = i
            break
    for i in range(location_section_start, min(location_section_start + 10, len(lines))):
        if lines[i] in ["United States", "USA", "U.S.", "U.S.A."]:
            city_line = i - 1 if i > 0 else 0
            return f"{lines[city_line]}, {lines[i]}"
    city_state_pattern = re.search(r"([A-Za-z\s]+),?\s+([A-Za-z\s]+)", text[:500])
    if city_state_pattern:
        return city_state_pattern.group(0)
    return "Location Not Found"

def extract_skills(text):
    skill_keywords = [
        "Python", "Flask", "Django", "Machine Learning", "SQL", "Java", "React", "AWS",
        "JavaScript", "HTML", "CSS", "TensorFlow", "Pandas", "NumPy", "Docker", "Kubernetes",
        "Git", "Azure", "Linux", "Node.js", "C#", "C++", "Go", "PHP", "TypeScript", "Tableau", 
        "Power BI", "Jupyter", "Spark", "Hadoop", "Scala", "Cloud", "Vagrant", "LLMs", "GPT",
        "Re-enforcement", "Site Reliability", "DevOps", "Microservices", "NOSQL", 
        "Apache Kafka", "Apache Webserver", "Blockchain", "Performance engineering",
        "AI model Training", "Data Science", "Feature Engineering", "AI", "Shell Script", 
        "Intrusion Detection", "Matlab", "R", "Agile", "SDLC"
    ]
    skills_section_match = re.search(r"Skills\n(.*?)(?:\n\n|\nProfile|\nEmployment)", text, re.DOTALL)
    skills_text = skills_section_match.group(1) if skills_section_match else text
    found_skills = [skill for skill in skill_keywords if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
    additional_skills = [line.strip() for line in skills_text.split("\n") if line.strip() and len(line.strip()) > 2]
    all_skills = list(set(found_skills + additional_skills))
    return all_skills

def extract_languages(text):
    languages_section = re.search(r"Languages\n(.*?)(?:\n\n|\nEducation|\nEmployment)", text, re.DOTALL)
    if languages_section:
        languages_text = languages_section.group(1)
        languages = [lang.strip() for lang in languages_text.split("\n") if lang.strip()]
        return languages
    common_languages = ["English", "Spanish", "French", "German", "Chinese", "Japanese", 
                        "Arabic", "Hindi", "Bengali", "Russian", "Portuguese"]
    found_languages = [lang for lang in common_languages if re.search(rf"\b{lang}\b", text)]
    return found_languages if found_languages else ["Not Found"]

def extract_experience(text):
    years_of_experience = re.findall(r"(\d{1,2})\+?\s*years?(?:\s+of\s+experience)?", text)
    if years_of_experience:
        max_experience = max(map(int, years_of_experience))
        return f"{max_experience} years of experience"
    employment_pattern = re.findall(r"(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", text, re.IGNORECASE)
    if employment_pattern:
        return f"{len(employment_pattern)} employment periods found"
    return "Experience not found"

def extract_qualifications(text):
    qualification_keywords = [
        "Bachelor's", "Master's", "PhD", "Degree", "Certification", "Diploma", "BSc", "MSc", 
        "B.A.", "M.A.", "BTech", "MTech", "MBA", "Engineering", "Architecture", "Computer Science", 
        "Information Technology", "Data Science", "Machine Learning", "AI", "Software Engineering",
        "CSSA", "Certified", "Business", "MIS", "University"
    ]
    education_section = re.search(r"Education\n(.*?)(?:\n\n|$)", text, re.DOTALL)
    if education_section:
        edu_text = education_section.group(1)
        education_entries = [line.strip() for line in edu_text.split("\n") if line.strip() 
                            and any(keyword in line for keyword in ["Bachelor", "Master", "PhD", "Degree"])]
        if education_entries:
            return education_entries
    found_qualifications = [qual for qual in qualification_keywords if re.search(rf"\b{qual}\b", text, re.IGNORECASE)]
    return list(set(found_qualifications)) if found_qualifications else ["Not Found"]

def extract_employment_history(text):
    employment_section = re.search(r"Employment History\n(.*?)(?:\nEducation|\n\nEducation|$)", text, re.DOTALL | re.IGNORECASE)
    if not employment_section:
        return ["Employment history not found"]
    employment_text = employment_section.group(1)
    job_pattern = re.findall(r"(.*?),\s+(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
                             employment_text, re.IGNORECASE | re.MULTILINE)
    if not job_pattern:
        job_pattern = re.findall(r"(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
                                employment_text, re.IGNORECASE | re.MULTILINE)
    if not job_pattern:
        lines = employment_text.split("\n")
        jobs = []
        for i, line in enumerate(lines):
            if re.search(r"\b(architect|analyst|lead|principal|manager|director|engineer)\b", line, re.IGNORECASE):
                if i < len(lines) - 1 and re.search(r"\d{4}", lines[i+1]):
                    jobs.append(line)
        return jobs if jobs else ["Could not parse employment details"]
    
    formatted_jobs = []
    for job in job_pattern:
        if len(job) == 5:
            formatted_jobs.append(f"{job[0]} at {job[1]}, {job[2]} ({job[3]} - {job[4]})")
        elif len(job) == 4:
            formatted_jobs.append(f"{job[0]} at {job[1]} ({job[2]} - {job[3]})")
    return formatted_jobs

def extract_profile_summary(text):
    profile_section = re.search(r"Profile\n(.*?)(?:\nEmployment|\n\nEmployment)", text, re.DOTALL)
    if profile_section:
        profile_text = profile_section.group(1).strip()
        if len(profile_text) > 500:
            return profile_text[:497] + "..."
        return profile_text
    alt_section = re.search(r"(Summary|About)\n(.*?)(?:\n\n|\nSkills|\nExperience)", text, re.DOTALL)
    if alt_section:
        return alt_section.group(2).strip()
    return "Profile summary not found"

# ---------------- Job Scraping and Matching Functions ----------------

def get_chrome_binary():
    possible_paths = [
        "/opt/google/chrome/chrome",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser"
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    return None

def init_chrome_driver():
    options = ChromeOptions()
    options.add_argument("--headless")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--remote-debugging-port=9222")
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64)")
    chrome_binary = get_chrome_binary()
    if not chrome_binary:
        raise Exception("No Chrome or Chromium binary found!")
    options.binary_location = chrome_binary
    if "google/chrome" in chrome_binary:
        driver_version = "133.0.6943.126"
    else:
        driver_version = "120.0.6099.224"
    driver_path = ChromeDriverManager(driver_version=driver_version).install()
    os.chmod(driver_path, 0o755)
    service = ChromeService(driver_path)
    driver = webdriver.Chrome(service=service, options=options)
    return driver

def init_firefox_driver():
    options = FirefoxOptions()
    options.add_argument("--headless")
    driver_path = GeckoDriverManager().install()
    service = FirefoxService(driver_path)
    driver = webdriver.Firefox(service=service, options=options)
    return driver

def init_driver():
    try:
        driver = init_chrome_driver()
        print("Initialized Chrome/Chromium driver.")
        return driver
    except Exception as e:
        print("Chrome driver initialization failed, trying Firefox. Error:", e)
        driver = init_firefox_driver()
        print("Initialized Firefox driver.")
        return driver

def scrape_dice(job_role, driver, preferred_location="", preferred_salary=""):
    print(f"🔍 Scraping Dice for: {job_role}")
    formatted_role = job_role.replace(" ", "%20")
    # Use location and salary from user preferences if provided
    loc_param = preferred_location if preferred_location and preferred_location != "Select Location" else ""
    salary_param = preferred_salary if preferred_salary and preferred_salary != "Select Salary Range" else ""
    url = f"https://www.dice.com/jobs?q={formatted_role}&location={loc_param}&salary={salary_param}&countryCode=US&page=1&pageSize=20&language=en"
    driver.get(url)
    time.sleep(5)
    soup = BeautifulSoup(driver.page_source, "html.parser")
    job_elements = soup.find_all("div", class_="card")
    jobs = []
    for job in job_elements:
        title_element = job.find("h5")
        link_element = job.find("a", href=True)
        location_element = job.find("span", attrs={"data-cy": "search-result-location"})
        company_name = job.find("a", attrs={"aria-label": lambda x: x and "company page for" in x})
        job_type_element = job.find("span", attrs={"data-cy": "search-result-employment-type"})
        posted_time_element = job.find("span", attrs={"data-cy": "card-posted-date"})
        updated_time_element = job.find("span", attrs={"data-cy": "card-modified-date"})
        description_element = job.find("div", attrs={"data-cy": "card-summary"})
        if title_element and link_element:
            jobs.append({
                "Job Role": job_role,
                "Title": title_element.text.strip(),
                "Apply Link": link_element["href"],
                "Location": location_element.text.strip() if location_element else "N/A",
                "Job Type": job_type_element.text.strip() if job_type_element else "N/A",
                "Posted Time": posted_time_element.text.strip() if posted_time_element else "N/A",
                "Updated Time": updated_time_element.text.strip() if updated_time_element else "N/A",
                "Description": description_element.text.strip() if description_element else "N/A",
                "Company Name": company_name.text.strip() if company_name else "N/A"
            })
    return jobs

def extract_career_goal_keywords(career_goals):
    if not career_goals:
        return {}
    goals_data = {
        "job_titles": [],
        "locations": [],
        "work_preferences": [],
        "industries": []
    }
    job_title_patterns = [
        r"(?:seeking|looking for|interested in)(?:\s+a)?\s+(?:role|position|job|career)?\s+(?:as|in)?\s+(?:an?)?\s+((?:senior|junior|lead|principal|staff)?\s*[a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))",
        r"(?:senior|junior|lead|principal|staff)?\s*([a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))\s+(?:role|position|job)",
        r"(?:become|be)\s+(?:a|an)\s+((?:senior|junior|lead|principal|staff)?\s*[a-z\s]+(?:engineer|developer|architect|manager|analyst|scientist|designer))"
    ]
    for pattern in job_title_patterns:
        matches = re.findall(pattern, career_goals.lower())
        if matches:
            goals_data["job_titles"].extend([title.strip().title() for title in matches])
    common_titles = [
        "software engineer", "data scientist", "machine learning engineer", 
        "devops engineer", "full stack developer", "frontend developer", 
        "backend developer", "site reliability engineer", "cloud engineer",
        "ai engineer", "web developer", "product manager", "project manager",
        "ux designer", "systems engineer", "solutions architect",
        "data engineer", "ml engineer", "research scientist"
    ]
    for title in common_titles:
        if title in career_goals.lower():
            goals_data["job_titles"].append(title.title())
    location_patterns = [
        r"(?:in|near|around|based in)\s+([a-z\s]+(?:area|city|region|county))",
        r"(?:in|near|around|based in)\s+([a-z]+,\s*[a-z]{2})",
        r"(?:in|near|around|based in)\s+([a-z\s]+)"
    ]
    for pattern in location_patterns:
        matches = re.findall(pattern, career_goals.lower())
        if matches:
            goals_data["locations"].extend([loc.strip().title() for loc in matches])
    common_cities = ["new york", "san francisco", "seattle", "boston", "austin", "chicago", 
                     "los angeles", "denver", "washington dc", "atlanta", "dallas"]
    for city in common_cities:
        if city in career_goals.lower():
            goals_data["locations"].append(city.title())
    preferences = ["remote", "hybrid", "onsite", "in-office", "work from home", "flexible", "part-time", "full-time"]
    for pref in preferences:
        if pref in career_goals.lower():
            goals_data["work_preferences"].append(pref.title())
    industries = ["tech", "healthcare", "finance", "fintech", "edtech", "e-commerce", "retail", 
                 "gaming", "entertainment", "education", "manufacturing", "automotive", 
                 "energy", "sustainability", "non-profit", "government", "consulting"]
    for industry in industries:
        if industry in career_goals.lower():
            goals_data["industries"].append(industry.title())
    return goals_data

def get_relevant_job_titles(parsed_data):
    """Determine job titles based on career goals and resume details.
       Priority is given if the career goal keywords appear in the resume skills."""
    job_titles = []
    career_goals = parsed_data.get('career_goals', "")
    career_keywords = extract_career_goal_keywords(career_goals)
    career_job_titles = career_keywords.get("job_titles", [])
    # Check if any career title keyword appears in the parsed skills
    valid_titles = []
    for title in career_job_titles:
        for skill in parsed_data.get('skills', []):
            if title.lower() in skill.lower():
                valid_titles.append(title)
                break
    if valid_titles:
        job_titles.extend(valid_titles)
    else:
        # If not, use career goals job titles (first priority) even if not in skills
        if career_job_titles:
            job_titles.extend(career_job_titles)
    # Fall back to employment history if needed
    if len(job_titles) < 3:
        for job in parsed_data.get('employment_history', []):
            if "at" in job:
                title = job.split("at")[0].strip().rstrip(",")
                job_titles.append(title)
        key_skills = []
        priority_skills = ["AI", "Machine Learning", "Python", "Java", "Cloud", "DevOps", "Architect", 
                           "Data Science", "Full Stack", "Backend", "Frontend", "Software Engineer"]
        for skill in priority_skills:
            if any(skill.lower() in s.lower() for s in parsed_data.get('skills', [])):
                key_skills.append(skill)
        for skill in key_skills[:3]:
            job_titles.append(f"{skill} Engineer")
            job_titles.append(f"{skill} Developer")
    unique_titles = list(set(job_titles))
    return unique_titles[:5]

def calculate_job_match_score(job, parsed_data, motivation_factors):
    """Calculate a job match score based on resume data and motivation factors."""
    total_score = 0
    max_possible_score = sum([factor['weight'] for factor in motivation_factors])
    factor_scores = {}
    for factor in motivation_factors:
        factor_name = factor['name']
        factor_weight = factor['weight']
        score = 0
        if factor_name == "Salary":
            salary_keywords = ["competitive salary", "excellent compensation", "high pay", "top pay"]
            if any(keyword in job['Description'].lower() for keyword in salary_keywords):
                score = factor_weight
            else:
                score = factor_weight * 0.5
        elif factor_name == "Location":
            # Use preferred_location if provided; otherwise fall back to parsed resume location
            user_location = parsed_data.get('preferred_location', parsed_data.get('location', '')).lower()
            if user_location and user_location != "location not found":
                if user_location in job['Location'].lower():
                    score = factor_weight
                elif "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
                    score = factor_weight * 0.8
                else:
                    score = 0
            else:
                score = factor_weight * 0.5
        elif factor_name == "Company Culture":
            culture_keywords = ["work-life balance", "inclusive", "diversity", "collaborative", 
                              "flexible", "innovative", "team-oriented", "growth opportunities"]
            matches = sum(1 for keyword in culture_keywords if keyword in job['Description'].lower())
            score = factor_weight * min(matches / 3, 1.0)
        elif factor_name == "Work-Life Balance":
            balance_keywords = ["flexible hours", "work-life balance", "remote work", "flexible schedule",
                              "family friendly", "paid time off", "pto", "vacation"]
            matches = sum(1 for keyword in balance_keywords if keyword in job['Description'].lower())
            score = factor_weight * min(matches / 2, 1.0)
        elif factor_name == "Growth Opportunities":
            growth_keywords = ["career advancement", "professional development", "promotion", "training",
                             "learning", "mentorship", "growth", "career path"]
            matches = sum(1 for keyword in growth_keywords if keyword in job['Description'].lower())
            score = factor_weight * min(matches / 2, 1.0)
        elif factor_name == "Job Stability":
            stability_keywords = ["established company", "long-term", "stable", "security", 
                                "permanent position", "industry leader"]
            matches = sum(1 for keyword in stability_keywords if keyword in job['Description'].lower())
            score = factor_weight * min(matches / 2, 1.0)
        elif factor_name == "Remote Work":
            if "remote" in job['Job Type'].lower() or "remote" in job['Description'].lower():
                score = factor_weight
            elif "hybrid" in job['Job Type'].lower() or "hybrid" in job['Description'].lower():
                score = factor_weight * 0.7
            else:
                score = 0
        elif factor_name == "Company Size":
            score = factor_weight * 0.5
        elif factor_name == "Industry Match":
            if 'career_goals' in parsed_data and parsed_data['career_goals']:
                career_goal_keywords = extract_career_goal_keywords(parsed_data['career_goals'])
                industries = career_goal_keywords.get("industries", [])
                if industries:
                    if any(industry.lower() in job['Description'].lower() for industry in industries):
                        score = factor_weight
                    else:
                        score = 0
                else:
                    score = factor_weight * 0.5
            else:
                score = factor_weight * 0.5
        elif factor_name == "Skill Match":
            if parsed_data.get('skills'):
                matches = sum(1 for skill in parsed_data['skills'] if skill.lower() in job['Description'].lower())
                score = factor_weight * min(matches / max(5, len(parsed_data['skills']) * 0.3), 1.0)
            else:
                score = factor_weight * 0.5
        else:
            score = factor_weight * 0.5
        factor_scores[factor_name] = score
        total_score += score
    percentage_score = (total_score / max_possible_score) * 100 if max_possible_score > 0 else 50
    if percentage_score >= 85:
        match_category = "Excellent"
    elif percentage_score >= 70:
        match_category = "Good"
    elif percentage_score >= 50:
        match_category = "Moderate"
    else:
        match_category = "Low"
    return {
        "total_score": total_score,
        "percentage_score": percentage_score,
        "match_category": match_category,
        "factor_scores": factor_scores
    }

# ---------------- Main Streamlit Application ----------------

def main():
    st.title("Enhanced Resume Parser with Job Matching")
    st.write("Upload your resume to extract details and find matching job opportunities.")
    
    if 'parsed_data' not in st.session_state:
        st.session_state.parsed_data = None
    if 'career_goals' not in st.session_state:
        st.session_state.career_goals = ""
    if 'preferred_location' not in st.session_state:
        st.session_state.preferred_location = "Select Location"
    if 'preferred_salary' not in st.session_state:
        st.session_state.preferred_salary = "Select Salary Range"
    if 'motivation_factors' not in st.session_state:
        st.session_state.motivation_factors = [
            {"name": "Salary", "weight": 5},
            {"name": "Location", "weight": 4},
            {"name": "Company Culture", "weight": 3},
            {"name": "Work-Life Balance", "weight": 3},
            {"name": "Growth Opportunities", "weight": 4},
            {"name": "Skill Match", "weight": 5}
        ]
    if 'job_results' not in st.session_state:
        st.session_state.job_results = None
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1

    # -------- Step 1: Upload and Parse Resume --------
    if st.session_state.current_step == 1:
        st.header("Step 1: Upload Your Resume")
        uploaded_file = st.file_uploader("Choose a PDF or DOCX resume", type=["pdf", "docx"])
        if uploaded_file is not None and st.button("Parse Resume", key="parse_resume_btn"):
            with st.spinner("Parsing Resume..."):
                try:
                    parsed_data = parse_resume(uploaded_file)
                    st.session_state.parsed_data = parsed_data
                    st.success("Resume parsed successfully!")
                    
                    # Enhanced formatted display of parsed resume details
                    st.subheader("Parsed Resume Details")
                    st.markdown("**Personal Information:**")
                    st.write(f"**Name:** {parsed_data['name']}")
                    st.write(f"**Email:** {parsed_data['email']}")
                    st.write(f"**Phone:** {parsed_data['phone']}")
                    st.write(f"**Location:** {parsed_data['location']}")
                    st.markdown("---")
                    st.markdown("**Experience & Skills:**")
                    st.write(f"**Experience:** {parsed_data['experience']}")
                    st.write(f"**Languages:** {', '.join(parsed_data['languages'])}")
                    st.write("**Skills:**")
                    st.write(", ".join(parsed_data['skills']))
                    st.markdown("---")
                    st.markdown("**Employment History:**")
                    for job in parsed_data['employment_history']:
                        st.write(f"- {job}")
                    st.markdown("---")
                    st.markdown("**Qualifications:**")
                    for qual in parsed_data['qualifications']:
                        st.write(f"- {qual}")
                    if parsed_data['profile_summary'] != "Profile summary not found":
                        st.markdown("---")
                        st.markdown("**Profile Summary:**")
                        st.write(parsed_data['profile_summary'])
                    
                    st.session_state.current_step = 2
                    st.experimental_rerun()
                except Exception as e:
                    st.error(f"Error parsing resume: {e}")

    # -------- Step 2: Career Purpose and Preferences --------
    elif st.session_state.current_step == 2:
        st.header("Step 2: Define Your Career Purpose and Preferences")
        st.write("Describe your career goals, desired job roles, and work preferences. Also select your preferred job location and desired salary range.")
        career_goals = st.text_area(
            "Career Goals and Preferences",
            value=st.session_state.career_goals,
            height=200,
            placeholder="Example: I am seeking a Senior Machine Learning Engineer role in the healthcare industry, preferably remote or in the Boston area. I'm interested in companies with a strong focus on AI ethics and work-life balance."
        )
        location = st.selectbox("Select Preferred Job Location", 
                                 ["Select Location", "New York", "San Francisco", "Los Angeles", "Boston", "Chicago", "Remote"],
                                 index=0)
        salary_range = st.selectbox("Select Desired Salary Range", 
                                 ["Select Salary Range", "Below $50k", "$50k-$70k", "$70k-$90k", "$90k-$110k", "Above $110k"],
                                 index=0)
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("← Back", key="back_to_step1"):
                st.session_state.current_step = 1
                st.experimental_rerun()
        with col2:
            if st.button("Save and Continue →", key="save_goals_btn"):
                if not career_goals.strip():
                    st.error("Career goals and preferences cannot be empty!")
                elif location == "Select Location":
                    st.error("Please select a job location.")
                elif salary_range == "Select Salary Range":
                    st.error("Please select a desired salary range.")
                else:
                    st.session_state.career_goals = career_goals
                    st.session_state.preferred_location = location
                    st.session_state.preferred_salary = salary_range
                    if st.session_state.parsed_data:
                        st.session_state.parsed_data['career_goals'] = career_goals
                        st.session_state.parsed_data['preferred_location'] = location
                        st.session_state.parsed_data['preferred_salary'] = salary_range
                    st.success("Career goals and preferences saved!")
                    
                    # Display extracted keywords from career goals
                    goals_keywords = extract_career_goal_keywords(career_goals)
                    if goals_keywords:
                        st.subheader("Extracted Keywords")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write("**Desired Job Titles:**")
                            for title in goals_keywords.get('job_titles', []):
                                st.write(f"- {title}")
                            st.write("**Preferred Locations:**")
                            for loc in goals_keywords.get('locations', []):
                                st.write(f"- {loc}")
                        with col2:
                            st.write("**Work Preferences:**")
                            for preference in goals_keywords.get('work_preferences', []):
                                st.write(f"- {preference}")
                            st.write("**Target Industries:**")
                            for industry in goals_keywords.get('industries', []):
                                st.write(f"- {industry}")
                    st.session_state.current_step = 3
                    st.experimental_rerun()
    
    # -------- Step 3: Motivation Matrix --------
    elif st.session_state.current_step == 3:
        st.header("Step 3: Customize Your Motivation Matrix")
        st.write("Adjust the importance of each factor in your job search by setting weights from 1-10.")
        updated_factors = []
        for factor in st.session_state.motivation_factors:
            weight = st.slider(
                f"{factor['name']}",
                min_value=1,
                max_value=10,
                value=factor['weight'],
                help=f"Set importance of {factor['name']} in your job search"
            )
            updated_factors.append({"name": factor['name'], "weight": weight})
        st.subheader("Add Custom Motivation Factor (Optional)")
        custom_factor = st.text_input("Factor Name")
        custom_weight = st.slider("Factor Weight", min_value=1, max_value=10, value=5)
        if st.button("Add Custom Factor", key="add_factor_btn") and custom_factor:
            updated_factors.append({"name": custom_factor, "weight": custom_weight})
            st.success(f"Added custom factor: {custom_factor}")
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("← Back", key="back_to_step2"):
                st.session_state.current_step = 2
                st.experimental_rerun()
        with col2:
            if st.button("Save and Continue →", key="save_matrix_btn"):
                st.session_state.motivation_factors = updated_factors
                st.success("Motivation matrix saved!")
                st.session_state.current_step = 4
                st.experimental_rerun()
    
    # -------- Step 4: Job Matching and Ranking --------
    elif st.session_state.current_step == 4:
        st.header("Step 4: Job Matching Results")
        if not st.session_state.parsed_data:
            st.warning("Please go back and upload your resume first.")
            if st.button("← Back to Step 1"):
                st.session_state.current_step = 1
                st.experimental_rerun()
        else:
            col1, col2 = st.columns([1, 3])
            with col1:
                if st.button("← Back", key="back_to_step3"):
                    st.session_state.current_step = 3
                    st.experimental_rerun()
            with col2:
                if st.button("Find Matching Jobs", key="find_jobs_btn"):
                    with st.spinner("Searching for matching jobs..."):
                        try:
                            job_titles = get_relevant_job_titles(st.session_state.parsed_data)
                            st.write("**Searching for the following job roles:**")
                            for title in job_titles:
                                st.write(f"- {title}")
                            driver = init_driver()
                            all_jobs = []
                            for job_title in job_titles:
                                jobs = scrape_dice(job_title, driver, st.session_state.preferred_location, st.session_state.preferred_salary)
                                all_jobs.extend(jobs)
                            driver.quit()
                            scored_jobs = []
                            for job in all_jobs:
                                match_score = calculate_job_match_score(
                                    job, 
                                    st.session_state.parsed_data,
                                    st.session_state.motivation_factors
                                )
                                job_with_score = {**job, "match_score": match_score}
                                scored_jobs.append(job_with_score)
                            sorted_jobs = sorted(
                                scored_jobs, 
                                key=lambda x: x["match_score"]["percentage_score"], 
                                reverse=True
                            )
                            st.session_state.job_results = sorted_jobs
                            st.success(f"Found {len(sorted_jobs)} matching jobs!")
                        except Exception as e:
                            st.error(f"Error finding matching jobs: {e}")
            
            if st.session_state.job_results:
                # ---------------- Summary Report (Step 5) ----------------
                st.subheader("Summary Report")
                summary = {"Excellent": 0, "Good": 0, "Moderate": 0, "Low": 0}
                for job in st.session_state.job_results:
                    cat = job["match_score"]["match_category"]
                    summary[cat] = summary.get(cat, 0) + 1
                st.write(summary)
                
                st.subheader("Top Matching Jobs")
                filter_col1, filter_col2 = st.columns(2)
                with filter_col1:
                    min_score = st.slider("Minimum Match Score (%)", min_value=0, max_value=100, value=50)
                with filter_col2:
                    match_categories = ["All"] + list(set([job["match_score"]["match_category"] for job in st.session_state.job_results]))
                    selected_category = st.selectbox("Match Category", match_categories)
                filtered_jobs = st.session_state.job_results
                if min_score > 0:
                    filtered_jobs = [job for job in filtered_jobs if job["match_score"]["percentage_score"] >= min_score]
                if selected_category != "All":
                    filtered_jobs = [job for job in filtered_jobs if job["match_score"]["match_category"] == selected_category]
                if not filtered_jobs:
                    st.warning("No jobs match the selected filters.")
                else:
                    jobs_df = pd.DataFrame([
                        {
                            "Title": job["Title"],
                            "Company": job["Company Name"],
                            "Location": job["Location"],
                            "Job Type": job["Job Type"],
                            "Match Score": f"{job['match_score']['percentage_score']:.1f}%",
                            "Match Category": job["match_score"]["match_category"],
                            "Posted": job["Posted Time"],
                            "Apply Link": job["Apply Link"]
                        }
                        for job in filtered_jobs
                    ])
                    st.download_button(
                        "Download Results as CSV",
                        jobs_df.to_csv(index=False).encode('utf-8'),
                        "job_matches.csv",
                        "text/csv",
                        key='download-csv'
                    )
                    for i, job in enumerate(filtered_jobs[:100]):
                        with st.expander(f"{job['Title']} at {job['Company Name']} - {job['match_score']['percentage_score']:.1f}% Match ({job['match_score']['match_category']})"):
                            job_col1, job_col2 = st.columns([2, 1])
                            with job_col1:
                                st.write(f"**Company:** {job['Company Name']}")
                                st.write(f"**Location:** {job['Location']}")
                                st.write(f"**Job Type:** {job['Job Type']}")
                                st.write(f"**Posted:** {job['Posted Time']}")
                                st.write("**Description:**")
                                st.write(job['Description'])
                                st.markdown(f"[Apply Now]({job['Apply Link']})")
                            with job_col2:
                                st.write("**Match Breakdown:**")
                                for factor, score in job["match_score"]["factor_scores"].items():
                                    max_score = next((f["weight"] for f in st.session_state.motivation_factors if f["name"] == factor), 5)
                                    st.write(f"- {factor}: {score}/{max_score}")
                                match_color = {
                                    "Excellent": "green",
                                    "Good": "blue",
                                    "Moderate": "orange",
                                    "Low": "red"
                                }.get(job["match_score"]["match_category"], "gray")
                                st.markdown(f"<div style='background-color:{match_color};height:20px;width:{job['match_score']['percentage_score']}%;border-radius:5px;'></div>", unsafe_allow_html=True)
                                st.write(f"{job['match_score']['percentage_score']:.1f}% Match")

if __name__ == "__main__":
    main()
